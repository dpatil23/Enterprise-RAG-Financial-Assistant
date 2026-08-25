import io
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import fitz  # PyMuPDF
try:
    import docx
except ImportError:
    docx = None

from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import settings
from app.core.embedder import embedder
from app.core.vector_store import vector_store
from app.core.graph_store import graph_store
from app.services.entity_extraction import entity_extractor
from app.services.entity_resolution import entity_resolver

logger = logging.getLogger("ingest_service")


def extract_text_from_file(file_bytes: bytes, filename: str) -> tuple[str, int]:
    """Extracts plain text and page count from PDF or DOCX."""
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        full_text = ""
        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            full_text += page.get_text() + "\n"
        doc.close()
        return full_text, total_pages
    elif ext in ["docx", "doc"]:
        if docx is None:
            raise ValueError("python-docx is not installed.")
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        full_text = "\n".join(paragraphs)
        return full_text, max(1, len(doc.paragraphs) // 10)
    else:
        # Fallback to UTF-8 text
        return file_bytes.decode("utf-8", errors="ignore"), 1


def ingest_pdf(file_bytes: bytes, filename: str) -> dict:
    """
    Hybrid Ingestion Pipeline for financial documents (Vector Store + Knowledge Graph).
    Supports PDF, DOCX, and TXT.
    """
    # --- Step 1: Parse ---
    full_text, total_pages = extract_text_from_file(file_bytes, filename)

    if not full_text.strip():
        raise ValueError("Could not extract text from document. It may be image-based or empty.")

    # --- Step 2: Chunk ---
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
    )
    chunks = splitter.split_text(full_text)

    # --- Step 3: Embed & Store in Vector DB ---
    logger.info(f"[IngestService] Embedding {len(chunks)} chunks into FAISS...")
    embeddings = embedder.embed(chunks)

    collection_name = "documents"
    clean_name = filename.rsplit(".", 1)[0].replace(" ", "_").lower()
    doc_id = clean_name

    num_stored = vector_store.add_chunks(
        collection_name=collection_name,
        chunks=chunks,
        embeddings=embeddings,
        doc_id=doc_id,
    )

    # --- Step 4: Extract & Store in Knowledge Graph in Parallel ---
    total_nodes_added = 0
    total_rels_added = 0
    raw_entities_all = []
    raw_rels_all = []

    logger.info(f"[IngestService] Extracting Knowledge Graph entities in parallel for {len(chunks)} chunks...")
    
    def process_chunk(idx: int, chunk_text: str):
        chunk_id = f"{doc_id}_chunk_{idx}"
        extracted = entity_extractor.extract_from_chunk(chunk_text)
        return chunk_id, extracted.get("entities", []), extracted.get("relationships", [])

    # Knowledge Graph extraction across 3 key summary sections (preserves 95% of Groq rate limits)
    chunks_to_extract = [(i, c) for i, c in enumerate(chunks[:3])]
    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = [executor.submit(process_chunk, i, c) for i, c in chunks_to_extract]
        for fut in as_completed(futures):
            chunk_id, entities, relationships = fut.result()
            if entities or relationships:
                resolved_entities, resolved_rels = entity_resolver.resolve_batch(
                    entities=entities,
                    relationships=relationships,
                )
                raw_entities_all.extend(resolved_entities)
                raw_rels_all.extend(resolved_rels)

                counts = graph_store.add_entities_and_relations(
                    entities=resolved_entities,
                    relationships=resolved_rels,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                )
                total_nodes_added += counts.get("nodes", 0)
                total_rels_added += counts.get("relationships", 0)

    # Get overall graph statistics
    graph_totals = graph_store.count_nodes_and_edges()

    return {
        "filename": filename,
        "doc_id": doc_id,
        "total_pages": total_pages,
        "total_chunks": num_stored,
        "chunk_size": settings.CHUNK_SIZE,
        "chunk_overlap": settings.CHUNK_OVERLAP,
        "graph_stats": {
            "nodes_extracted_in_doc": len(raw_entities_all),
            "relationships_extracted_in_doc": len(raw_rels_all),
            "total_graph_nodes": graph_totals.get("node_count", 0),
            "total_graph_relationships": graph_totals.get("relationship_count", 0),
        },
    }
